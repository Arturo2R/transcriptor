import requests
import streamlit as st
import pandas as pd
import math
from pathlib import Path
from pydub import AudioSegment
import os
import io
import subprocess
from openai import OpenAI
from docx import Document
import functools
# Set the title and favicon that appear in the Browser's tab bar.
st.set_page_config(
    page_title='Transcripción de Audio',
    # page_icon=':pen:', # This is an emoji shortcode. Could be a URL too.
)

# -----------------------------------------------------------------------------
# Declare some useful functions.


def split_audio_file(input_file, max_size_mb=20 * 1024 * 1024):
   # Load the audio file
    audio = AudioSegment.from_file(input_file, format="mp3")
    
    # Get the file extension

    ext = "mp3"
    
    # Calculate the number of chunks needed
    file_size = input_file.size
    num_chunks = math.ceil(file_size / (max_size_mb * 1024 * 1024))
    
    # Get the duration of the audio file in seconds
    audio_duration_seconds = audio.duration_seconds
    
    # Convert duration to minutes
    audio_duration_minutes:int = math.ceil(audio_duration_seconds / 60)
    
    # Calculate the duration of each chunk
    chunk_length_ms = len(audio) // num_chunks
    files = []
    # Split and export the audio chunks
    for i in range(num_chunks):
        start_time = i * chunk_length_ms
        end_time = (i + 1) * chunk_length_ms if i < num_chunks - 1 else len(audio)
        
        chunk = audio[start_time:end_time]
        output_filename = f"{input_file.name}_part{i+1}.{ext}"
        chunk.export(output_filename, format=ext)
        files.append(output_filename)
        print(f"Exported {output_filename}")

    print(f"Audio file split into {num_chunks} parts.")
    return files, audio_duration_minutes

def get_trm() -> float:
    url = f"https://www.datos.gov.co/resource/32sa-8pi3.json"
    response = requests.get(url)
    data = response.json()
    trm:int = data[0]['valor']
    return trm


# -----------------------------------------------------------------------------
# Draw the actual page

# Set the title that appears at the top of the page.
st.markdown('''
# :earth_americas: Transcripción de Audio

Esta aplicación es una interfaz para utilizar el modelo Whisper-large V2 apartir de la API de OpenAI, el costo de transcripción es de 0.006 por minuto. Solo tienes que poner tu [apikey de openai](https://openai.com) y luego puedes seleccionar cualquier archivo para transcribirlo.

''')

# Add some spacing
''
''
loading = False
if 'done' not in st.session_state:
    st.session_state.done = False

if 'audio_chunks' not in st.session_state:
    st.session_state.audio_chunks = []

if 'cost' not in st.session_state:
    st.session_state.cost = 0


col1, col2 = st.columns(2)

with st.form(key='transcription_form', clear_on_submit=False, border=False):
    # st.mark
    with col1:
        key = st.text_input(label="Open AI ApiKey", placeholder="sk-proj-pTJxcSi3ofbeOO0JaZn8M33l32kFJVn8KnG9P44WamCkp6nMZ")
        # col11, col12 = st.columns(2)
        # with col11:
        language = st.selectbox("Idioma", ["es", "en"])
        # with col12:
        #     timestamps = st.checkbox("Timestamps")

    if key:
        client = OpenAI(api_key=key)

    with col2:
        # st.markdown("## Sube Tu Archivo")
        audio = st.file_uploader(type="mp3", label="Archivo De Audio")

    
    if audio:
        # trm = get_trm()
        cost:int = math.ceil(AudioSegment.from_file(audio).duration_seconds /60)
        st.markdown(f"La transcripción del audio costará aproximadamente  `${cost * 0.006}USD` ó `${cost * 0.006 * 4000}COP`")
        st.audio(audio)

    submitted = st.form_submit_button("Transcribir")
    
    # Saving the chunks to verify the output
if submitted and audio:
    with st.spinner('Diviendo el archivo...'):
        audio_chunks, audio_duration = split_audio_file(audio)
    with st.spinner('Leyendo Archivos...'):
        for chunk in audio_chunks:
            audio_file = open(chunk, "rb")
            audio_bytes = audio_file.read()

    if key:
        with st.spinner('Transcribiendo...'):
            transcript = client.audio.transcriptions.create(
                file=audio_file,
                model="whisper-1",
                response_format="text",
                language=language or "en",
                timestamp_granularities=["segment"],
                )
        st.text("✅ Transcritó correctamente") 
        st.markdown(transcript)

        # Create a Word document with the transcript content

        # Create a new Word document
        doc = Document()

        # Add the transcript content to the document
        # for result in transcript["results"]:
        #     if timestamps:
        #         # Add the timestamp and transcript text as a paragraph
        #         doc.add_paragraph(f"{result['alternatives'][0]['timestamps'][0][0]} - {result['alternatives'][0]['transcript']}")
        #     else:
        #         # Add only the transcript text as a paragraph
        #         doc.add_paragraph(result['alternatives'][0]['transcript'])
        doc.add_paragraph(transcript)
        # Save the Word document
        col112, col122 = st.columns(2)
        doc.save(f"{audio.name}_transcripción.docx")
        with open(f'{audio.name}_transcripción.docx', 'rb') as f:
            with col112:
                st.download_button('Descargar DOCX', f, file_name=f'{audio.name}_transcripción.docx', mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with open(f'{audio.name}_transcripción.txt', 'w') as file:
            file.write(transcript)
        with open(f'{audio.name}_transcripción.txt', 'rb') as file:
            with col122:
                st.download_button('Descargar TXT',file, file_name=f'{audio.name}_transcripción.txt', mime="text/plain")
        os.remove(chunk)
    else:
        st.error("Por favor, ingresa tu apikey de OpenAI")



    



# st.checkbox("Markdown")

# st.button("Transcribir", on_click=)

st.sidebar.markdown('''

_Creado por Arturo Rebolledo Econometrista y Desarrollador de Software_
''')

